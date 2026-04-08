from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_DIR = Path("reports") / "bdlm_union_first_pass"
DOCX_PATH = OUTPUT_DIR / "union_submission_polished.docx"
FIGURE_PATH = OUTPUT_DIR / "union_abstract_figure_confirmed_signal.png"


TITLE = (
    "Low confirmed post-treatment failure after a pretomanid-free "
    "bedaquiline-delamanid-linezolid strategy with or without moxifloxacin "
    "for rifampicin-resistant tuberculosis in China"
)

ABSTRACT = {
    "Background": (
        "Pretomanid-containing short-course regimens have reshaped treatment expectations "
        "for rifampicin-resistant tuberculosis (RR-TB), but pretomanid remains unavailable "
        "in some settings. We evaluated post-treatment outcomes of a pretomanid-free strategy "
        "based on bedaquiline, delamanid and linezolid, with or without moxifloxacin (BDL/M), "
        "in an ongoing strategy-guided clinical study in China."
    ),
    "Methods": (
        "We conducted an interim descriptive analysis of protocol-consistent participants managed "
        "with BDL or BDLM. Because composite unfavorable outcomes may overstate biological treatment "
        "failure when follow-up attrition is substantial, we reconstructed strict post-treatment "
        "endpoints that separated confirmed failure/relapse from loss to follow-up and withdrawal. "
        "For this analysis, the primary endpoint was confirmed failure/relapse by 6 months after "
        "treatment completion; the secondary endpoint was confirmed failure/relapse by 12 months. "
        "Month-2 culture conversion and end-of-treatment culture status were also summarized."
    ),
    "Results": (
        "Among 122 entered records, 121 unique participants remained after deduplication; 108 "
        "protocol-consistent participants were included (49 BDL, 59 BDLM). Among evaluable "
        "participants, month-2 culture conversion was 90/101 (89.1%; 95% CI 81.3-94.4), and "
        "end-of-treatment culture negativity was 78/78 (100.0%; 95% CI 95.4-100.0). Confirmed "
        "failure/relapse occurred in 2/108 participants by 6 months after treatment completion "
        "(1.9%; 95% CI 0.2-6.5) and remained 2/108 by 12 months (1.9%; 95% CI 0.2-6.5). At 6 months, "
        "17/108 participants (15.7%) were classified as lost to follow-up and 8/108 (7.4%) as withdrawn; "
        "at 12 months, these proportions were 36/108 (33.3%) and 8/108 (7.4%), respectively. Both "
        "confirmed failure/relapse events occurred in the BDL group; none were observed in the BDLM group, "
        "although treatment allocation was strategy-guided rather than randomized."
    ),
    "Conclusions": (
        "In this interim analysis from a pretomanid-unavailable setting, a pretomanid-free BDL/M "
        "strategy showed strong early microbiological response and a low rate of confirmed post-treatment "
        "failure/relapse. Apparent long-term unfavorable outcomes were driven primarily by follow-up "
        "attrition rather than confirmed biological recurrence. Participants classified as lost to "
        "follow-up at 12 months remain under active tracing."
    ),
}

FIGURE_TITLE = "Key efficacy signals of a pretomanid-free BDL/M strategy for RR-TB in China"
FIGURE_LEGEND = (
    "Figure. Key efficacy signals of a pretomanid-free BDL/M strategy for rifampicin-resistant "
    "tuberculosis in China. Left panel shows month-2 culture conversion and end-of-treatment culture "
    "negativity; right panel shows confirmed failure/relapse by 6 and 12 months after treatment "
    "completion. Points represent proportions and whiskers exact 95% confidence intervals. Loss to "
    "follow-up and withdrawal are described in the text; participants classified as lost to follow-up "
    "at 12 months remain under active tracing."
)


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(11)


def add_section_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.space_after = Pt(3)


def add_abstract_paragraph(document: Document, heading: str, body: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run_head = p.add_run(f"{heading}: ")
    run_head.bold = True
    run_body = p.add_run(body)
    run_body.bold = False


def approximate_word_count(text: str) -> int:
    return len(text.split())


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_default_font(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(TITLE)
    title_run.bold = True
    title_run.font.size = Pt(15)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run("Preferred category: TBScience abstract").italic = True

    word_count = approximate_word_count(" ".join(ABSTRACT.values()))
    wc_p = doc.add_paragraph()
    wc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    wc_p.add_run(f"Approximate abstract word count: {word_count}").italic = True

    add_section_heading(doc, "Polished Abstract")
    for heading, body in ABSTRACT.items():
        add_abstract_paragraph(doc, heading, body)

    doc.add_paragraph()
    add_section_heading(doc, "Preferred Figure")

    fig_title_p = doc.add_paragraph()
    fig_title_run = fig_title_p.add_run(FIGURE_TITLE)
    fig_title_run.bold = True

    if FIGURE_PATH.exists():
        fig_p = doc.add_paragraph()
        fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig_p.add_run().add_picture(str(FIGURE_PATH), width=Inches(6.7))

    fig_leg_p = doc.add_paragraph()
    fig_leg_p.paragraph_format.space_after = Pt(6)
    fig_leg_p.add_run(FIGURE_LEGEND)

    doc.add_paragraph()
    notes_head = doc.add_paragraph()
    notes_head.add_run("Submission Notes").bold = True
    for note in [
        "Keep the main efficacy phrase as 'confirmed failure/relapse'.",
        "Avoid causal claims between BDL and BDLM because treatment allocation was strategy-guided rather than randomized.",
        "If the abstract system has a strict word cap, shorten the Methods paragraph first.",
    ]:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        p.add_run(note)

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


if __name__ == "__main__":
    main()
