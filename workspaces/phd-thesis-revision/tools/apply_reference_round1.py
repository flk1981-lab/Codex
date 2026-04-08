from __future__ import annotations

import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent
FINAL_DIR = ROOT / "output" / "final"
REPORT_PATH = ROOT / "output" / "references" / "reference_round1_report.json"


MAIN_NEW_REFERENCES = [
    "[102] Fu L, Xiong J, Wang H, et al. Study protocol for safety and efficacy of all-oral shortened regimens for multidrug-resistant tuberculosis: a multicenter randomized withdrawal trial and a single-arm trial [SEAL-MDR][J]. BMC Infect Dis, 2023, 23(1): 834.",
    "[103] World Health Organization. WHO consolidated guidelines on tuberculosis: module 4: treatment and care[R]. Geneva: World Health Organization, 2025.",
    "[104] World Health Organization. WHO operational handbook on tuberculosis: module 4: treatment and care[R]. Geneva: World Health Organization, 2025.",
]

REVIEW_NEW_REFERENCES = [
    "[28] IVANOVA O, HOFFMANN V S, LANGE C, et al. Post-tuberculosis lung impairment: systematic review and meta-analysis of spirometry data from 14 621 people[J]. Eur Respir Rev, 2023, 32(168): 220221.",
    "[29] WALLIS R S, HAFNER R. Advancing host-directed therapy for tuberculosis[J]. Nat Rev Immunol, 2015, 15(4): 255-263.",
    "[30] NDONG SIMA C A A, SMITH D, PETERSEN D C, et al. The immunogenetics of tuberculosis (TB) susceptibility[J]. Immunogenetics, 2023, 75(3): 215-230.",
]

TEXT_UPDATES = [
    {
        "label": "main-platform-citation-gap",
        "needle": "[103-105]",
        "replace": ("[103-105]", "[102-104]"),
    },
    {
        "label": "main-ae-qtc",
        "needle": "QTcF ≥500 ms事件仅见于Bdq组",
        "citation": "[7,21,84]",
    },
    {
        "label": "main-mechanism-support",
        "needle": "从生物学解释角度看，xCT/SLC7A11轴仍是理解本章结果的重要理论支点。",
        "citation": "[30,56,59]",
    },
    {
        "label": "main-conclusion-signal",
        "needle": "临床部分最稳妥的结论是：在MDR-Chin平台内，含SASP方案在pre-XDR-TB患者中显示出积极但仍属探索性的疗效与安全性信号。",
        "citation": "[17,22,60]",
    },
    {
        "label": "main-conclusion-boundary",
        "needle": "但这部分证据仍来自前瞻性、非随机、非盲法队列研究",
        "citation": "[17,22,60]",
    },
    {
        "label": "review-intro-burden",
        "needle": "结核病（tuberculosis, TB）仍是全球重要的传染病死因之一。",
        "citation": "[1,28,29]",
    },
    {
        "label": "review-xct-node",
        "needle": "当前证据提示，xCT/SLC7A11 并非单纯感染相关标志物",
        "citation": "[7,8]",
    },
    {
        "label": "review-integrated-framework",
        "needle": "这一整合框架的最大价值，在于它把SASP从“一个老药的偶然再利用”提升为“一个具有理论闭环的HDT候选策略”。",
        "citation": "[9,11,29]",
    },
    {
        "label": "review-immunogenetics",
        "needle": "这一机制链条的重要性在于，它使“基因型差异”不再只是一个黑箱统计现象",
        "citation": "[8,30]",
    },
    {
        "label": "review-precision-pathway",
        "needle": "针对结核病中xCT/SLC7A11–SASP轴提出的精准宿主导向治疗路径。",
        "citation": "[25-27]",
    },
    {
        "label": "review-endpoints",
        "needle": "对xCT/SASP路径而言，候选临床终点也不能只盯着培养转阴。",
        "citation": "[4,26,28]",
    },
    {
        "label": "review-overall-summary",
        "needle": "总体而言，xCT/SLC7A11轴代表了一类值得在TB中持续推进的宿主靶点",
        "citation": "[12,29,30]",
    },
]


def append_citation(text: str, citation: str) -> str:
    if citation in text:
        return text
    for punct in ("。", "；", ".", "！", "？"):
        if text.endswith(punct):
            return f"{text[:-1]}{citation}{punct}"
    return f"{text}{citation}"


def find_paragraph(doc: Document, needle: str):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    return None


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_references_before(doc: Document, anchor_text: str, entries: list[str]) -> int:
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == anchor_text:
            anchor = paragraph
            break
    if anchor is None:
        raise ValueError(f"Anchor paragraph not found: {anchor_text}")

    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in entries:
            delete_paragraph(paragraph)

    inserted = 0
    for entry in entries:
        new_paragraph = anchor.insert_paragraph_before(entry)
        if anchor.style is not None:
            new_paragraph.style = anchor.style
        inserted += 1
    return inserted


def main() -> int:
    report: list[dict[str, object]] = []
    docs = sorted(FINAL_DIR.glob("*.docx"))
    if not docs:
        raise FileNotFoundError(f"No .docx files found in {FINAL_DIR}")

    for docx_path in docs:
        doc = Document(str(docx_path))
        updates_applied: list[str] = []

        for update in TEXT_UPDATES:
            paragraph = find_paragraph(doc, update["needle"])
            if paragraph is None and "replace" in update:
                _, new = update["replace"]
                paragraph = find_paragraph(doc, new)
            if paragraph is None:
                raise ValueError(f"Paragraph not found in {docx_path.name}: {update['needle']}")

            original = paragraph.text
            if "replace" in update:
                old, new = update["replace"]
                if old in paragraph.text:
                    paragraph.text = paragraph.text.replace(old, new)
            if "citation" in update:
                paragraph.text = append_citation(paragraph.text, update["citation"])
            if paragraph.text != original:
                updates_applied.append(update["label"])

        main_inserted = insert_references_before(doc, "文献综述", MAIN_NEW_REFERENCES)
        review_inserted = insert_references_before(doc, "附录", REVIEW_NEW_REFERENCES)

        doc.save(str(docx_path))
        report.append(
            {
                "file": str(docx_path),
                "text_updates_applied": updates_applied,
                "main_references_inserted": main_inserted,
                "review_references_inserted": review_inserted,
            }
        )

    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
