from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent.parent
FINAL_DIR = ROOT / "output" / "final"
REPORT_PATH = ROOT / "output" / "references" / "reference_style_cleanup_report.json"
STRATEGY_PATH = ROOT / "output" / "references" / "citation_strategy.md"


REPLACEMENTS = {
    "[99] Rubin E J, Mizrahi V. Shortening the Short Course of Tuberculosis Treatment[J]. N Engl J Med, 2021,384(18):1764-1765.":
        "[99] Rubin E J, Mizrahi V. Shortening the Short Course of Tuberculosis Treatment[J]. N Engl J Med, 2021, 384(18): 1764-1765.",
    "[100] Dartois V, Rubin E J. Shortening Tuberculosis Treatment - A Strategic Retreat[J]. N Engl J Med, 2023,388(10):939-941.":
        "[100] Dartois V, Rubin E J. Shortening Tuberculosis Treatment - A Strategic Retreat[J]. N Engl J Med, 2023, 388(10): 939-941.",
    "[101] Dartois V A, Mizrahi V, Savic R M, et al. Strategies for shortening tuberculosis therapy[J]. Nat Med, 2025,31(6):1765-1775.":
        "[101] Dartois V A, Mizrahi V, Savic R M, et al. Strategies for shortening tuberculosis therapy[J]. Nat Med, 2025, 31(6): 1765-1775.",
}


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def append_strategy_note() -> None:
    note = """

## 9. 参考文献体例清理

- 统一修正主体参考文献 `[99]-[101]` 中缺失的卷期/页码空格。
- 删除综述参考文献列表中新增文献前遗留的空白段，保证列表连续。
"""
    text = STRATEGY_PATH.read_text(encoding="utf-8")
    if "## 9. 参考文献体例清理" not in text:
        STRATEGY_PATH.write_text(text.rstrip() + note + "\n", encoding="utf-8")


def main() -> int:
    report: list[dict[str, object]] = []
    for docx_path in sorted(FINAL_DIR.glob("*.docx")):
        doc = Document(str(docx_path))
        replaced = 0
        blanks_removed = 0

        for paragraph in doc.paragraphs:
            new_text = REPLACEMENTS.get(paragraph.text)
            if new_text:
                paragraph.text = new_text
                replaced += 1

        paragraphs = list(doc.paragraphs)
        for idx in range(1, len(paragraphs) - 1):
            prev_text = paragraphs[idx - 1].text.strip()
            this_text = paragraphs[idx].text.strip()
            next_text = paragraphs[idx + 1].text.strip()
            if prev_text.startswith("[27]") and this_text == "" and next_text.startswith("[28]"):
                delete_paragraph(paragraphs[idx])
                blanks_removed += 1
                break

        doc.save(str(docx_path))
        report.append(
            {
                "file": str(docx_path),
                "references_reformatted": replaced,
                "blank_paragraphs_removed": blanks_removed,
            }
        )

    append_strategy_note()
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
