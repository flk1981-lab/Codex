from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent.parent
FINAL_DIR = ROOT / "output" / "final"
REPORT_PATH = ROOT / "output" / "references" / "reference_round4_report.json"
STRATEGY_PATH = ROOT / "output" / "references" / "citation_strategy.md"


INLINE_UPDATES = [
    (
        "在耐药肺结核中，HDT的潜在价值尤为突出。",
        "[24,25,29]",
    ),
    (
        "目前，结核病HDT研究涉及抗炎药、免疫调节剂、代谢调节剂、维生素及多种老药新用策略。",
        "[24-29,39,40]",
    ),
    (
        "围绕SASP在结核病中的应用价值，前期工作已积累了临床和实验两方面线索。",
        "[22,60]",
    ),
    (
        "在实验层面，前期研究和基金设计提示，SASP抑制xCT后，可增强巨噬细胞对结核菌的清除能力，并减轻感染诱导的过度炎症反应。",
        "[30,49,50,56,59,92]",
    ),
    (
        "尽管SASP在总体上显示出应用潜力，但并不意味着所有患者都可能获得同等程度获益。",
        "[14,30,42,56,68]",
    ),
    (
        "综合全文证据，当前最稳妥的总体判断是：SASP在耐药肺结核中已经形成“值得继续推进”的临床与实验联合证据基础，但尚未达到可以提出确定性临床推荐的程度。",
        "[17,22,30,60]",
    ),
]


FUTURE_OUTLOOK_PARAGRAPHS = [
    "未来研究最值得优先推进的方向，不是简单重复“再比较一次是否加用SASP”，而是把前几章已经暴露出的关键问题真正连成可验证的转化链条：在更严格的前瞻性研究中同步完成宿主分层、药物暴露检测、动态机制读出和多维临床终点评估。只有这样，才能区分“药物本身无效”“靶点未被命中”以及“只在特定宿主亚群中有效”这几种完全不同的情形[17,24,29,30]。",
    "从临床转化角度看，后续研究还应把长期毒性和长期功能结局放到与近期病原学结局同等重要的位置。对于耐药肺结核尤其是pre-XDR-TB患者而言，真正有价值的策略不仅要帮助提高治疗成功率，还应尽可能减少Lzd相关长期神经毒性、降低治疗后肺功能损害和生活质量下降的代价。因此，若未来SASP相关研究能够在“疗效、机制、分层和长期结局”四个层面同时取得一致证据，其临床意义将明显超出单一辅助加药策略本身[22,28,105-107]。",
]


def find_paragraph(doc: Document, needle: str):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    return None


def append_citation(text: str, citation: str) -> str:
    if citation in text:
        return text
    trailing = re.search(r"(\[[0-9,\- ]+\])([。；.!！？]?)$", text)
    if trailing:
        existing = trailing.group(1)
        punct = trailing.group(2)
        current_nums: list[int] = []
        for part in re.split(r"[, ]+", existing.strip("[]")):
            if not part:
                continue
            if "-" in part:
                a, b = part.split("-", 1)
                current_nums.extend([int(a), int(b)])
            else:
                current_nums.append(int(part))
        new_nums: list[int] = []
        for part in re.split(r"[, ]+", citation.strip("[]")):
            if not part:
                continue
            if "-" in part:
                a, b = part.split("-", 1)
                new_nums.extend([int(a), int(b)])
            else:
                new_nums.append(int(part))
        merged = sorted(set(current_nums + new_nums))
        merged_text = "[" + ",".join(str(n) for n in merged) + "]"
        return text[: trailing.start(1)] + merged_text + punct
    for punct in ("。", "；", ".", "！", "？"):
        if text.endswith(punct):
            return f"{text[:-1]}{citation}{punct}"
    return f"{text}{citation}"


def collapse_adjacent_citations(text: str) -> str:
    pattern = re.compile(r"\[([0-9,\- ]+)\]\[([0-9,\- ]+)\]")
    while True:
        match = pattern.search(text)
        if not match:
            return text
        nums: list[int] = []
        for raw in (match.group(1), match.group(2)):
            for part in re.split(r"[, ]+", raw):
                if not part:
                    continue
                if "-" in part:
                    a, b = part.split("-", 1)
                    nums.extend([int(a), int(b)])
                else:
                    nums.append(int(part))
        merged = "[" + ",".join(str(n) for n in sorted(set(nums))) + "]"
        text = text[: match.start()] + merged + text[match.end() :]


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    if paragraph.style is not None:
        new_para.style = paragraph.style
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def delete_exact_text(doc: Document, text: str) -> None:
    paragraph = find_paragraph(doc, text)
    if paragraph is not None and paragraph.text == text:
        delete_paragraph(paragraph)


def ensure_future_outlook(doc: Document) -> int:
    anchor = find_paragraph(doc, "7.5.2 未来展望")
    if anchor is None:
        raise ValueError("Future outlook heading not found")

    for text in FUTURE_OUTLOOK_PARAGRAPHS:
        delete_exact_text(doc, text)
    inserted = 0
    for text in reversed(FUTURE_OUTLOOK_PARAGRAPHS):
        insert_paragraph_after(anchor, text)
        inserted += 1
    return inserted


def append_strategy_note() -> None:
    note = """

## 8. 第 4 轮新增处理

### 第一章前言收口补引

- 对第一章中 5 个仍偏概括的关键论断段补入更贴切的 1-3 篇文献：
  - HDT 在耐药结核中的潜在价值
  - HDT 研究谱系与局限
  - SASP 的临床前与临床转化基础
  - 宿主遗传异质性与分层应用

### 第七章未来展望收口

- 在 `7.5.2 未来展望` 下新增 2 段可直接用于答辩口径的展望性总结。
- 重点统一三点：
  - 未来研究要做“证据链桥接”，不是重复现象学比较
  - 长期毒性与长期功能结局应作为核心终点
  - SASP 的价值取决于疗效、机制、分层和长期结局四条线是否同时成立
"""
    text = STRATEGY_PATH.read_text(encoding="utf-8")
    if "## 8. 第 4 轮新增处理" not in text:
        STRATEGY_PATH.write_text(text.rstrip() + note + "\n", encoding="utf-8")


def main() -> int:
    report: list[dict[str, object]] = []
    for docx_path in sorted(FINAL_DIR.glob("*.docx")):
        doc = Document(str(docx_path))
        inline_updates: list[str] = []
        for needle, citation in INLINE_UPDATES:
            paragraph = find_paragraph(doc, needle)
            if paragraph is None:
                raise ValueError(f"Paragraph not found in {docx_path.name}: {needle}")
            updated = append_citation(paragraph.text, citation)
            updated = collapse_adjacent_citations(updated)
            if updated != paragraph.text:
                paragraph.text = updated
                inline_updates.append(citation)

        inserted = ensure_future_outlook(doc)
        for paragraph in doc.paragraphs:
            collapsed = collapse_adjacent_citations(paragraph.text)
            if collapsed != paragraph.text:
                paragraph.text = collapsed
        doc.save(str(docx_path))
        report.append(
            {
                "file": str(docx_path),
                "inline_updates": inline_updates,
                "future_outlook_blocks_inserted": inserted,
            }
        )

    append_strategy_note()
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
