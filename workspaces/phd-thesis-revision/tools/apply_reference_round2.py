from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent.parent
FINAL_DIR = ROOT / "output" / "final"
REPORT_PATH = ROOT / "output" / "references" / "reference_round2_report.json"
STRATEGY_PATH = ROOT / "output" / "references" / "citation_strategy.md"


MAIN_NEW_REFERENCES = [
    "[105] Lan Z, Ahmad N, Baghaei P, et al. Drug-associated adverse events in the treatment of multidrug-resistant tuberculosis: an individual patient data meta-analysis[J]. Lancet Respir Med, 2020, 8(4): 383-394.",
    "[106] Wasserman S, Brust J C M, Abdelwahab M T, et al. Linezolid toxicity in patients with drug-resistant tuberculosis: a prospective cohort study[J]. J Antimicrob Chemother, 2022, 77(4): 1146-1154.",
    "[107] Ivanova O, Hoffmann V S, Lange C, et al. Post-tuberculosis lung impairment: systematic review and meta-analysis of spirometry data from 14 621 people[J]. Eur Respir Rev, 2023, 32(168): 220221.",
]

CH4_PARAGRAPHS = [
    "本研究在停药后36个月仍观察到持续性的Lzd相关神经或视觉症状，提示这类毒性并不一定会随停药迅速消退。既往针对DR-TB患者的临床观察也提示，Lzd相关周围神经病和视神经病可在治疗结束后持续存在，且部分患者恢复并不完全[9,10,87]。",
    "从治疗策略角度看，这一发现与更大样本的MDR-TB不良事件分析及Lzd前瞻性毒性队列研究方向一致：Lzd虽显著提升了耐药结核全口服方案的活性，但也始终是导致严重不良事件、剂量调整或永久停药的重要药物之一，因此未来短程/中短程方案优化必须同步考虑长期功能结局，而不能只看近期培养学获益[105-107]。",
]

REVIEW_APPEND_CITATIONS = [
    (
        "围绕 xCT/SLC7A11–SASP 轴的精准宿主导向治疗，不太可能依赖单一标志物完成分层和决策，而更需要整合基线宿主分层、药物暴露评估、动态机制读出和多维临床终点。",
        "[25,26]",
    ),
    (
        "需要强调的是，这些动态标志物的价值，不在于把它们变成新的“诊断TB工具”，而在于让它们服务于治疗决策。",
        "[25,26]",
    ),
    (
        "再次，患者中心终点不应被忽视。",
        "[28]",
    ),
]


def append_citation(text: str, citation: str) -> str:
    if citation in text:
        return text
    for punct in ("。", "；", ".", "！", "？"):
        if text.endswith(punct):
            return f"{text[:-1]}{citation}{punct}"
    return f"{text}{citation}"


def find_paragraph(doc: Document, needle: str):
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    return None


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    if paragraph.style is not None:
        new_para.style = paragraph.style
    return new_para


def insert_references_before(doc: Document, anchor_text: str, entries: list[str]) -> int:
    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == anchor_text:
            anchor = paragraph
            break
    if anchor is None:
        raise ValueError(f"Anchor not found: {anchor_text}")

    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in entries:
            delete_paragraph(paragraph)

    inserted = 0
    for entry in entries:
        new_paragraph = anchor.insert_paragraph_before(entry)
        if anchor.style is not None:
            new_paragraph.style = anchor.style
        inserted += 1
    return inserted


def ensure_ch4_insertions(doc: Document) -> int:
    heading = find_paragraph(doc, "4.4.2 利奈唑胺长期毒性的临床启示")
    if heading is None:
        raise ValueError("Chapter 4 heading not found")

    inserted = 0
    for text in CH4_PARAGRAPHS:
        existing = find_paragraph(doc, text)
        if existing is not None:
            heading = existing
            continue
        heading = insert_paragraph_after(heading, text)
        inserted += 1
    return inserted


def patch_review_citations(doc: Document) -> list[str]:
    changed: list[str] = []
    for needle, citation in REVIEW_APPEND_CITATIONS:
        paragraph = find_paragraph(doc, needle)
        if paragraph is None:
            raise ValueError(f"Review paragraph not found: {needle}")
        updated = append_citation(paragraph.text, citation)
        if updated != paragraph.text:
            paragraph.text = updated
            changed.append(citation)
    return changed


def append_strategy_note() -> None:
    note = """

## 6. 第 2 轮新增处理

### 主体

- 新增 `[105]-[107]`，用于补强第四章关于 `Lzd` 长期毒性、严重不良事件负担和长期功能结局的论证。
- 在 `4.4.2 利奈唑胺长期毒性的临床启示` 下新增 2 段说明性讨论：
  - 一段强调神经/视觉毒性可能持续存在
  - 一段强调方案优化不能只看近期培养学结局，还要兼顾长期功能结局

### 综述

- 在精准路径部分进一步补入已有综述内编号文献：
  - 多维标志物/监测框架 `[25,26]`
  - 动态标志物服务治疗决策 `[25,26]`
  - 患者中心终点 `[28]`
"""
    text = STRATEGY_PATH.read_text(encoding="utf-8")
    if "## 6. 第 2 轮新增处理" not in text:
        STRATEGY_PATH.write_text(text.rstrip() + note + "\n", encoding="utf-8")


def main() -> int:
    report: list[dict[str, object]] = []
    for docx_path in sorted(FINAL_DIR.glob("*.docx")):
        doc = Document(str(docx_path))
        main_refs_inserted = insert_references_before(doc, "文献综述", MAIN_NEW_REFERENCES)
        ch4_insertions = ensure_ch4_insertions(doc)
        review_changes = patch_review_citations(doc)
        doc.save(str(docx_path))
        report.append(
            {
                "file": str(docx_path),
                "main_references_inserted": main_refs_inserted,
                "chapter4_paragraphs_ensured": ch4_insertions,
                "review_citation_updates": review_changes,
            }
        )

    append_strategy_note()
    REPORT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
